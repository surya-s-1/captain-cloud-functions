import os
import io
import json
import docx
import mimetypes
import pandas as pd
import functions_framework
from urllib.parse import urlparse
from google.cloud import storage, firestore, documentai_v1 as documentai

# --- Environment Variables ---

GOOGLE_CLOUD_PROJECT = os.environ.get('GOOGLE_CLOUD_PROJECT')
PROCESSOR_ID = os.getenv('DOC_AI_PROCESSOR_ID')
LOCATION = os.getenv('DOC_AI_LOCATION')
FIRESTORE_DATABASE = os.getenv('FIRESTORE_DATABASE')
OUTPUT_BUCKET = os.getenv('OUTPUT_BUCKET')

storage_client = storage.Client()
firestore_client = firestore.Client(database=FIRESTORE_DATABASE)
documentai_client = documentai.DocumentProcessorServiceClient()
PROCESSOR_NAME = documentai_client.processor_path(
    GOOGLE_CLOUD_PROJECT, LOCATION, PROCESSOR_ID
)

# --- Helper Functions ---


def _update_firestore_status(project_id, version, status):
    '''Updates the status of a specific project version in Firestore.'''
    doc_ref = (
        firestore_client.collection('projects')
        .document(project_id)
        .collection(f'versions')
        .document(version)
    )
    update_data = {'status': status}
    doc_ref.set(update_data, merge=True)


def _download_blob_to_memory(bucket_name, source_blob_name):
    '''Downloads a blob from a GCS bucket to an in-memory byte stream.'''
    bucket = storage_client.bucket(bucket_name)

    print(f'Attempting to download gs://{bucket_name}/{source_blob_name}')

    blob = bucket.blob(source_blob_name)
    file_stream = io.BytesIO()
    blob.download_to_file(file_stream)
    file_stream.seek(0)

    print(f'Successfully downloaded gs://{bucket_name}/{source_blob_name}')

    return file_stream


def _extract_from_structured(file_url):
    """Extracts text from a CSV or Excel file using pandas, with location data."""
    parsed_url = urlparse(file_url)
    bucket_name = parsed_url.netloc
    blob_name = parsed_url.path.lstrip('/')

    print(f'Extracting from structured file: gs://{bucket_name}/{blob_name}')

    file_stream = _download_blob_to_memory(bucket_name, blob_name)

    if blob_name.lower().endswith('.csv'):
        df = pd.read_csv(file_stream)
    elif blob_name.lower().endswith(('.xlsx', '.xls')):
        df = pd.read_excel(file_stream, engine='openpyxl')
    else:
        raise ValueError(f"Unsupported structured file type: {blob_name}")

    extracted_data = []

    for row_index, row in df.iterrows():
        extracted_data.append(
            {
                'text': row.to_string(),
                'location': {'row': int(row_index), 'column_headers': list(df.columns)},
            }
        )
    return {
        'file_type': 'structured',
        'extracted_by': 'pandas',
        'extracted_text': extracted_data,
    }


def _extract_from_word(file_url):
    """Extracts text from a Word document using python-docx, with location data."""
    parsed_url = urlparse(file_url)
    bucket_name = parsed_url.netloc
    blob_name = parsed_url.path.lstrip('/')

    print(f'Extracting from Word document: gs://{bucket_name}/{blob_name}')

    file_stream = _download_blob_to_memory(bucket_name, blob_name)
    doc = docx.Document(file_stream)
    extracted_data = []

    for para_index, para in enumerate(doc.paragraphs):
        if para.text.strip():
            extracted_data.append(
                {'text': para.text, 'location': {'paragraph_number': para_index}}
            )
    return {
        'file_type': 'unstructured',
        'extracted_by': 'python-docx',
        'extracted_text': extracted_data,
    }


def _extract_from_document_ai(file_url):
    """Extracts text from a PDF or image using Google Document AI OCR, with location data."""
    parsed_url = urlparse(file_url)
    bucket_name = parsed_url.netloc
    blob_name = parsed_url.path.lstrip('/')

    print(f'Extracting from Document AI supported file: gs://{bucket_name}/{blob_name}')
    mime_type, _ = mimetypes.guess_type(file_url)
    if not mime_type:
        raise ValueError(f"Could not determine MIME type for file: {file_url}")

    print(f'Downloading file for Document AI processing: {file_url}')
    file_stream = _download_blob_to_memory(bucket_name, blob_name)
    content = file_stream.read()

    raw_document = documentai.RawDocument(content=content, mime_type=mime_type)
    request = documentai.ProcessRequest(name=PROCESSOR_NAME, raw_document=raw_document)
    response = documentai_client.process_document(request=request)

    extracted_data = []
    for page_index, page in enumerate(response.document.pages):
        for paragraph in page.paragraphs:
            paragraph_text = ''
            for segment in paragraph.layout.text_anchor.text_segments:
                start_index = int(segment.start_index)
                end_index = int(segment.end_index)
                paragraph_text += response.document.text[start_index:end_index]

            if paragraph_text.strip():
                extracted_data.append(
                    {'text': paragraph_text, 'location': {'page': page_index + 1}}
                )
    return {
        'file_type': 'semistructured',
        'extracted_by': 'document_ai',
        'extracted_text': extracted_data,
    }


# --- Main Cloud Function (HTTP Trigger) ---
@functions_framework.http
def process_documents(request):
    '''
    Cloud Function to extract text from files specified in an HTTP POST request.
    It processes the request synchronously and returns the results upon completion.
    '''
    try:
        message_payload = request.get_json(silent=True)
        if not message_payload:
            return (
                json.dumps(
                    {
                        'status': 'error',
                        'message': 'No JSON payload found in request body.',
                    }
                ),
                400,
            )

        project_id = message_payload.get('project_id', None)
        version = message_payload.get('version', None)
        file_urls = message_payload.get('files', [])

        if not project_id or not version or not file_urls:
            return (
                json.dumps(
                    {
                        'status': 'error',
                        'message': 'Required details (project_id, version, or files) are missing.',
                    }
                ),
                400,
            )

        _update_firestore_status(project_id, version, 'START_TEXT_EXTRACT')

        extracted_results = []

        for file_url in file_urls:
            try:
                file_name = os.path.basename(urlparse(file_url).path)
                file_extension = file_name.split('.')[-1].lower()

                result_object = {'file_name': file_name, 'file_url': file_url}

                if file_extension in ['csv', 'xlsx', 'xls']:
                    print(
                        f'Detected structured file type: {file_extension}. Using pandas.'
                    )
                    result_object.update(_extract_from_structured(file_url))

                elif file_extension in ['docx']:
                    print(f'Detected Word document. Using python-docx.')
                    result_object.update(_extract_from_word(file_url))

                elif file_extension in ['pdf', 'jpg', 'jpeg', 'png']:
                    print(
                        f'Detected Document AI supported file type: {file_extension}. Using Document AI.'
                    )
                    result_object.update(_extract_from_document_ai(file_url))

                else:
                    print(f'Skipping unsupported file type: {file_url}')
                    continue

                extracted_results.append(result_object)
            except Exception as e:
                print(f'Error processing file {file_url}: {e}')

        if not extracted_results:
            raise Exception('Nothing extracted.')

        output_blob_path = f'extracted-text/{project_id}/v_{version}.json'

        print(
            f'All files processed. Uploading results to: {OUTPUT_BUCKET}/{output_blob_path}'
        )

        output_bucket = storage_client.bucket(OUTPUT_BUCKET)
        output_blob = output_bucket.blob(output_blob_path)
        output_blob.upload_from_string(
            data=json.dumps(extracted_results, indent=2),
            content_type='application/json',
        )

        extracted_text_url = f'gs://{OUTPUT_BUCKET}/{output_blob_path}'
        print(f'Successfully wrote results to {extracted_text_url}')

        _update_firestore_status(project_id, version, 'COMPLETE_TEXT_EXTRACT')

        return (
            json.dumps(
                {
                    'status': 'success',
                    'project_id': project_id,
                    'version': version,
                    'extracted_text_url': extracted_text_url,
                }
            ),
            200,
        )

    except Exception as e:
        print(f'An error occurred: {e}')

        _update_firestore_status(project_id, version, 'ERR_TEXT_EXTRACT')

        return json.dumps({'status': 'error', 'message': str(e)}), 500
